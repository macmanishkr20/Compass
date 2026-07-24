"""Attachment ingestion — shared by Home/Chat and the Agent Console.

Turns raw uploaded files (base64 data URLs from the browser) into model-ready
content:

  * raster images  -> gpt-5 vision parts (passed through as image_url),
  * PDF / DOCX      -> extracted text, inlined,
  * ZIP archives    -> each text/code entry extracted and inlined,
  * text / code     -> decoded and inlined,
  * anything else   -> a short "[binary file]" placeholder.

Extraction is best-effort and never raises to the caller: a file that can't be
read is noted inline so the turn still runs. `build_user_message()` is the one
entry point both engines call, so images and files behave identically in Chat
and in the Agent Console.
"""

from __future__ import annotations

import base64
import io
import zipfile
from typing import Any

from compass.models.messages import Message, user_message

# Guards so a huge upload can't blow the model's context window.
MAX_TEXT_CHARS = 200_000  # per extracted file
MAX_ZIP_ENTRIES = 300
MAX_ZIP_TOTAL_CHARS = 400_000

# Extensions treated as text when pulled out of a ZIP (skip binaries/images).
_TEXT_EXTS = {
    "txt", "md", "markdown", "rst", "log", "csv", "tsv", "json", "yaml", "yml",
    "toml", "ini", "cfg", "conf", "env", "xml", "html", "htm", "css", "scss",
    "js", "jsx", "ts", "tsx", "mjs", "cjs", "vue", "svelte", "py", "pyi",
    "java", "kt", "kts", "c", "h", "cpp", "cc", "hpp", "cs", "go", "rs", "rb",
    "php", "swift", "m", "mm", "sh", "bash", "zsh", "sql", "graphql", "gql",
    "proto", "dockerfile", "gitignore", "makefile", "gradle", "properties",
    "svg", "tex", "r", "jl", "lua", "pl", "dart", "scala", "clj", "ex", "exs",
}


def _ext(name: str) -> str:
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def _decode_data_url(data_url: str) -> bytes:
    """Bytes from a `data:...;base64,XXXX` URL (or a bare base64 string)."""
    if not data_url:
        return b""
    payload = data_url.split(",", 1)[1] if data_url.startswith("data:") else data_url
    try:
        return base64.b64decode(payload)
    except Exception:  # noqa: BLE001
        return b""


def _clip(text: str, limit: int = MAX_TEXT_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n… [truncated, {len(text) - limit} more characters]"


def _as_text(data: bytes) -> str:
    """Decode bytes as text, or flag them as binary."""
    if b"\x00" in data[:8192]:
        return f"[binary file, {len(data)} bytes — not shown as text]"
    return _clip(data.decode("utf-8", errors="replace"))


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = "\n\n".join(p.strip() for p in pages if p.strip())
        return _clip(text) if text.strip() else "[PDF had no extractable text]"
    except Exception as err:  # noqa: BLE001
        return f"[could not read PDF: {err}]"


def _extract_docx(data: bytes) -> str:
    try:
        import docx

        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return _clip(text) if text.strip() else "[DOCX had no extractable text]"
    except Exception as err:  # noqa: BLE001
        return f"[could not read DOCX: {err}]"


def _extract_zip(data: bytes) -> str:
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except Exception as err:  # noqa: BLE001
        return f"[could not read ZIP: {err}]"
    out: list[str] = []
    total = 0
    for info in zf.infolist():
        if info.is_dir() or len(out) >= MAX_ZIP_ENTRIES:
            continue
        name = info.filename
        if _ext(name) not in _TEXT_EXTS and "." in name:
            continue  # skip binaries/images inside the archive
        try:
            raw = zf.read(info)
        except Exception:  # noqa: BLE001
            continue
        body = _as_text(raw)
        chunk = f"--- {name} ---\n{body}"
        total += len(chunk)
        out.append(chunk)
        if total >= MAX_ZIP_TOTAL_CHARS:
            out.append("… [archive truncated]")
            break
    if not out:
        return "[ZIP contained no readable text files]"
    return "\n\n".join(out)


def process_attachment(att: dict) -> dict | None:
    """Normalize one raw upload `{name, mime, data_url}` into either
    `{kind:'image', name, data_url}` or `{kind:'text', name, text}`."""
    name = (att.get("name") or "file").strip()
    mime = (att.get("mime") or "").lower()
    data_url = att.get("data_url") or ""
    # Some callers may pre-extract text (kept for compatibility).
    if att.get("text") is not None and not data_url:
        return {"kind": "text", "name": name, "text": _clip(str(att["text"]))}

    ext = _ext(name)
    if mime.startswith("image/") and mime != "image/svg+xml":
        return {"kind": "image", "name": name, "data_url": data_url}

    data = _decode_data_url(data_url)
    if ext == "pdf" or mime == "application/pdf":
        return {"kind": "text", "name": name, "text": _extract_pdf(data)}
    if ext == "docx" or "wordprocessingml" in mime:
        return {"kind": "text", "name": name, "text": _extract_docx(data)}
    if ext == "zip" or mime in ("application/zip", "application/x-zip-compressed"):
        return {"kind": "text", "name": name, "text": _extract_zip(data)}
    # svg + any text/code file
    return {"kind": "text", "name": name, "text": _as_text(data)}


def build_user_message(text: str, raw_attachments: list[dict] | None) -> Message:
    """The one message builder for both engines. Text attachments are inlined
    as fenced blocks; images become multimodal image_url parts for gpt-5
    vision. With no images the content stays a plain string (cheapest path)."""
    processed = [
        p for p in (process_attachment(a) for a in (raw_attachments or [])) if p
    ]
    text_files = [p for p in processed if p["kind"] == "text"]
    images = [p for p in processed if p["kind"] == "image" and p.get("data_url")]

    body = text or ""
    for f in text_files:
        body += f"\n\n--- Attached file: {f['name']} ---\n```\n{f.get('text', '')}\n```"

    if not images:
        return user_message(body)

    parts: list[dict[str, Any]] = [
        {"type": "text", "text": body or "(see the attached image)"}
    ]
    for img in images:
        parts.append({"type": "image_url", "image_url": {"url": img["data_url"]}})
    return Message(role="user", content=parts)  # type: ignore[arg-type]
